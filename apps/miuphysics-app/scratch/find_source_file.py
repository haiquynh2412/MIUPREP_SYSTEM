import os
import fitz
import docx

dir_path = r"C:\Users\HAIQUYNH\OneDrive\SACH VIET\VAT LY"
keyword = "BỒI DƯỠNG HỌC SINH GIỎI VẬT LÍ 9"

def check_pdf(path):
    try:
        doc = fitz.open(path)
        for i in range(len(doc)):
            text = doc[i].get_text()
            if keyword in text:
                return True, i + 1, len(doc)
    except Exception as e:
        pass
    return False, 0, 0

def check_docx(path):
    try:
        doc = docx.Document(path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        text = "\n".join(fullText)
        if keyword in text:
            return True, len(doc.paragraphs)
    except Exception as e:
        pass
    return False, 0

output_file = r"C:\Users\HAIQUYNH\OneDrive\CODE AI\MIUPREP_SYSTEM\apps\miuphysics-app\scratch\search_source_result.txt"

with open(output_file, "w", encoding="utf-8") as f:
    f.write(f"Searching for keyword: '{keyword}' in {dir_path}\n")
    for filename in os.listdir(dir_path):
        path = os.path.join(dir_path, filename)
        if filename.lower().endswith(".pdf"):
            found, page, total = check_pdf(path)
            if found:
                f.write(f"[FOUND] PDF: {filename} | Matches on Page: {page}/{total} | Path: {path}\n")
        elif filename.lower().endswith(".docx"):
            found, paras = check_docx(path)
            if found:
                f.write(f"[FOUND] DOCX: {filename} | Paragraphs: {paras} | Path: {path}\n")

print(f"Search completed! Results written to {output_file}")
