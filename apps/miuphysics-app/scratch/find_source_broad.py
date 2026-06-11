import os
import fitz
import docx

root_dir = r"C:\Users\HAIQUYNH\OneDrive\SACH VIET"
keyword = "BỒI DƯỠNG HỌC SINH GIỎI VẬT LÍ 9"

def check_pdf(path):
    try:
        doc = fitz.open(path)
        for i in range(len(doc)):
            text = doc[i].get_text()
            if keyword in text:
                return True, i + 1, len(doc)
    except Exception:
        pass
    return False, 0, 0

def check_docx(path):
    try:
        doc = docx.Document(path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        text = "\n".join(fullText)
        if keyword in text:
            return True, len(doc.paragraphs)
    except Exception:
        pass
    return False, 0

output_file = r"C:\Users\HAIQUYNH\OneDrive\CODE AI\MIUPREP_SYSTEM\apps\miuphysics-app\scratch\search_source_broad_result.txt"

with open(output_file, "w", encoding="utf-8") as f:
    f.write(f"Broad searching for keyword: '{keyword}' in {root_dir}\n")
    for root, dirs, files in os.walk(root_dir):
        # Skip zipped folders or large unrelated directories
        if "node_modules" in root or ".git" in root:
            continue
        for file in files:
            path = os.path.join(root, file)
            if file.lower().endswith(".pdf"):
                found, page, total = check_pdf(path)
                if found:
                    f.write(f"[FOUND] PDF: {file} | Matches on Page: {page}/{total} | Path: {path}\n")
                    print(f"Found PDF: {file}")
            elif file.lower().endswith(".docx"):
                found, paras = check_docx(path)
                if found:
                    f.write(f"[FOUND] DOCX: {file} | Paragraphs: {paras} | Path: {path}\n")
                    print(f"Found DOCX: {file}")

print(f"Search completed! Results written to {output_file}")
